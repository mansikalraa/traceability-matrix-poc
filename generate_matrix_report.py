#!/usr/bin/env python3
"""
Traceability Matrix & Documentation Generator Script

Automates generation of single document files (.docx or .pdf) from the repository's
Traceability Matrix and documentation markdown files.

Execution Modes:
  Mode 1 (--mode 1 / --summary-only):
    1. Cover Page with Traceability Matrix
    2. Incomplete Rows & Open Issues Audit Report

  Mode 2 (--mode 2 / --full) [Default]:
    Executes all 4 steps:
    1. Cover Page with Traceability Matrix
    2. Incomplete Rows & Open Issues Audit Report
    3. Stacked Markdown Documentation (Use Cases -> Risk Management -> Test Cases)
       with Latest PR Review Details (Author, Reviewers, PR Title, PR Description)
    4. Issue and Pull Request Linkage (Cover Page & Documentation linking to Issues and PRs)
"""

import argparse
import html
import json
import os
import re
import ssl
import subprocess
import sys
import urllib.parse
import urllib.request
from datetime import datetime
from xml.etree import ElementTree

# Word Document Generation
import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.shared import Inches, Pt, RGBColor

# PDF Generation
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from reportlab.platypus import (
    HRFlowable,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


# ==============================================================================
# GitHub API & Metadata Fetcher
# ==============================================================================

class GitHubClient:
    """Handles GitHub API calls to fetch issues, PRs, and review details."""

    def __init__(self, token=None, repo_owner="mansikalraa", repo_name="traceability-matrix-poc"):
        self.repo_owner = repo_owner
        self.repo_name = repo_name
        self.token = token or self._detect_token()
        self.ssl_ctx = ssl.create_default_context()
        self.ssl_ctx.check_hostname = False
        self.ssl_ctx.verify_mode = ssl.CERT_NONE

    def _detect_token(self):
        """Attempts to find GitHub token in env or git credentials."""
        token = os.environ.get("GITHUB_TOKEN") or os.environ.get("GH_TOKEN")
        if token:
            return token
        try:
            res = subprocess.run(
                ["git", "credential", "fill"],
                input="protocol=https\nhost=github.com\n",
                text=True,
                capture_output=True,
                timeout=5,
            )
            for line in res.stdout.splitlines():
                if line.startswith("password="):
                    return line.split("=", 1)[1]
        except Exception:
            pass
        return None

    def _make_request(self, endpoint):
        url = f"https://api.github.com/repos/{self.repo_owner}/{self.repo_name}/{endpoint}"
        headers = {"User-Agent": "TraceabilityMatrixScript/1.0", "Accept": "application/vnd.github.v3+json"}
        if self.token:
            headers["Authorization"] = f"token {self.token}"

        req = urllib.request.Request(url, headers=headers)
        try:
            with urllib.request.urlopen(req, context=self.ssl_ctx, timeout=10) as resp:
                return json_loads(resp.read().decode("utf-8"))
        except Exception as e:
            return None

    def fetch_all_issues(self):
        """Fetches all issues from repository."""
        data = self._make_request("issues?state=all&per_page=100")
        return data or []

    def fetch_all_prs(self):
        """Fetches all pull requests from repository."""
        data = self._make_request("pulls?state=all&per_page=100")
        return data or []

    def fetch_pr_reviews(self, pr_number):
        """Fetches reviews for a given PR."""
        data = self._make_request(f"pulls/{pr_number}/reviews")
        return data or []


def json_loads(data_str):
    import json
    return json.loads(data_str)


# ==============================================================================
# Matrix Parser & Audit Engine
# ==============================================================================

class TraceabilityMatrix:
    """Parses markdown matrix and computes status audit."""

    def __init__(self, readme_path="README.md"):
        self.readme_path = readme_path
        self.headers = []
        self.rows = []
        self._parse()

    def _parse(self):
        if not os.path.exists(self.readme_path):
            return
        with open(self.readme_path, "r", encoding="utf-8") as f:
            content = f.read()

        lines = content.splitlines()
        in_matrix = False
        table_lines = []

        for line in lines:
            if "## Traceability Matrix" in line or "| Requirement |" in line:
                in_matrix = True
            if in_matrix:
                if line.strip().startswith("|"):
                    table_lines.append(line.strip())
                elif table_lines and not line.strip().startswith("|"):
                    break

        if not table_lines:
            return

        # Parse headers
        header_line = table_lines[0]
        self.headers = [c.strip() for c in header_line.split("|")[1:-1]]

        # Skip delimiter row
        data_lines = [l for l in table_lines[1:] if not re.match(r"^\|?\s*:?-+:?\s*\|", l)]

        for line in data_lines:
            cells = [c.strip() for c in line.split("|")[1:-1]]
            row_data = []
            for cell in cells:
                cell_info = self._parse_cell(cell)
                row_data.append(cell_info)
            self.rows.append(row_data)

    def _parse_cell(self, cell_str):
        # Extract markdown link e.g. [FR-001](https://github.com/...)
        link_match = re.search(r"\[([^\]]+)\]\(([^)]+)\)", cell_str)
        if link_match:
            text = link_match.group(1).strip()
            url = link_match.group(2).strip()
            issue_match = re.search(r"/issues/(\d+)", url)
            issue_num = int(issue_match.group(1)) if issue_match else None
            return {"text": text, "url": url, "issue_number": issue_num, "is_link": True}
        else:
            text = cell_str.strip()
            return {"text": text, "url": None, "issue_number": None, "is_link": False}


def audit_matrix(matrix, issues_map, prs_list):
    """
    Evaluates incomplete rows in the traceability matrix.
    Row is incomplete if:
    1. Columns are not mapped (e.g. '--' or missing issue link)
    2. Linked issue is still OPEN.
    """
    audit_results = []

    for idx, row in enumerate(matrix.rows, start=1):
        reasons = []
        is_incomplete = False

        req_cell = row[0] if len(row) > 0 else {"text": "--"}
        hazard_cell = row[1] if len(row) > 1 else {"text": "--"}
        design_cell = row[2] if len(row) > 2 else {"text": "--"}
        tc_cell = row[3] if len(row) > 3 else {"text": "--"}
        verif_cell = row[4] if len(row) > 4 else {"text": "--"}

        # 1. Missing Column Mapping Checks
        if tc_cell["text"] in ["--", "", "N/A"]:
            is_incomplete = True
            reasons.append("Test Case not mapped ('--')")

        if hazard_cell["text"] in ["--", "", "N/A"]:
            is_incomplete = True
            reasons.append("Hazard not mapped ('--')")

        if verif_cell["text"] in ["--", "", "N/A"]:
            is_incomplete = True
            reasons.append("Verification status pending ('--')")

        # 2. Check Linked Issues Status
        for cell in [req_cell, tc_cell]:
            issue_num = cell.get("issue_number")
            if issue_num and issue_num in issues_map:
                issue_info = issues_map[issue_num]
                if issue_info.get("state") == "open":
                    is_incomplete = True
                    reasons.append(f"Issue #{issue_num} ({cell['text']}) is OPEN")

        # Find linked PRs
        linked_prs = []
        for cell in [req_cell, tc_cell]:
            code = cell["text"]
            issue_num = cell.get("issue_number")
            for pr in prs_list:
                pr_body = pr.get("body") or ""
                pr_title = pr.get("title") or ""
                if (
                    code != "--"
                    and (code in pr_title or code in pr_body)
                    or (issue_num and f"#{issue_num}" in pr_body)
                    or (issue_num and pr.get("number") == issue_num)
                ):
                    if pr not in linked_prs:
                        linked_prs.append(pr)

        audit_results.append({
            "row_index": idx,
            "req": req_cell,
            "hazard": hazard_cell,
            "design": design_cell,
            "tc": tc_cell,
            "verif": verif_cell,
            "is_incomplete": is_incomplete,
            "reasons": reasons,
            "linked_prs": linked_prs,
        })

    return audit_results


# ==============================================================================
# Markdown Document Parser
# ==============================================================================

def load_markdown_files(docs_dir="docs"):
    """
    Loads markdown files in sequence:
    1. Use cases (docs/use-cases/*.md)
    2. Risk management (docs/risk-management.md)
    3. Test cases (docs/test-cases/*.md)
    """
    sections = []

    # 1. Use Cases
    use_cases_dir = os.path.join(docs_dir, "use-cases")
    if os.path.exists(use_cases_dir):
        files = sorted(os.listdir(use_cases_dir))
        for fname in files:
            if fname.endswith(".md"):
                fpath = os.path.join(use_cases_dir, fname)
                sections.append({"category": "Use Case", "file": fname, "path": fpath})

    # 2. Risk Management
    risk_file = os.path.join(docs_dir, "risk-management.md")
    if os.path.exists(risk_file):
        sections.append({"category": "Risk Analysis", "file": "risk-management.md", "path": risk_file})

    # 3. Test Cases
    tc_dir = os.path.join(docs_dir, "test-cases")
    if os.path.exists(tc_dir):
        files = sorted(os.listdir(tc_dir))
        for fname in files:
            if fname.endswith(".md"):
                fpath = os.path.join(tc_dir, fname)
                sections.append({"category": "Test Case", "file": fname, "path": fpath})

    return sections


def parse_md_elements(filepath):
    """Simple parser to read markdown into structured elements."""
    if not os.path.exists(filepath):
        return []
    with open(filepath, "r", encoding="utf-8") as f:
        lines = f.readlines()

    elements = []
    in_code = False
    code_buf = []

    for line in lines:
        stripped = line.rstrip("\n")
        if stripped.startswith("```"):
            if in_code:
                elements.append({"type": "code", "text": "\n".join(code_buf)})
                code_buf = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_buf.append(stripped)
            continue

        if not stripped.strip():
            elements.append({"type": "spacer"})
            continue

        if stripped.startswith("# "):
            elements.append({"type": "h1", "text": stripped[2:].strip()})
        elif stripped.startswith("## "):
            elements.append({"type": "h2", "text": stripped[3:].strip()})
        elif stripped.startswith("### "):
            elements.append({"type": "h3", "text": stripped[4:].strip()})
        elif stripped.startswith("- ") or stripped.startswith("* "):
            elements.append({"type": "bullet", "text": stripped[2:].strip()})
        elif re.match(r"^\d+\.\s", stripped):
            item_text = re.sub(r"^\d+\.\s", "", stripped).strip()
            elements.append({"type": "numbered", "text": item_text})
        elif stripped == "---":
            elements.append({"type": "hr"})
        else:
            elements.append({"type": "paragraph", "text": stripped.strip()})

    return elements


# ==============================================================================
# DOCX Document Builder
# ==============================================================================

class DocxReportBuilder:
    """Builds formatted Word (.docx) documents."""

    def __init__(self, output_path="traceability_report.docx"):
        self.output_path = output_path
        self.doc = docx.Document()
        self._setup_document_styles()

    def _setup_document_styles(self):
        # Set 1 inch margins
        for section in self.doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Style colors
        self.PRIMARY_COLOR = RGBColor(30, 58, 138)  # Deep Navy #1E3A8A
        self.SECONDARY_COLOR = RGBColor(71, 85, 105)  # Slate Gray
        self.ALERT_COLOR = RGBColor(185, 28, 28)  # Crimson Red

    def _add_hyperlink(self, paragraph, url, text, color="1E40AF", underline=True):
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        hyperlink = parse_xml(
            f'<w:hyperlink xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            f' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
            f' r:id="{r_id}"/>'
        )
        new_run = parse_xml(f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        rPr = parse_xml(f'<w:rPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')

        if color:
            c = parse_xml(f'<w:color xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="{color}"/>')
            rPr.append(c)
        if underline:
            u = parse_xml(f'<w:u xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="single"/>')
            rPr.append(u)

        new_run.append(rPr)
        t = parse_xml(
            f'<w:t xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            f' xml:space="preserve">{html.escape(text)}</w:t>'
        )
        new_run.append(t)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    def _set_cell_background(self, cell, hex_color):
        shading = parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{hex_color}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    def build_report(self, matrix, audit_results, stacked_docs, gh_client, mode="2"):
        # ----------------------------------------------------------------------
        # STEP 1: COVER PAGE
        # ----------------------------------------------------------------------
        title_p = self.doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_p.add_run("TRACEABILITY MATRIX &\nCOMPLIANCE REPORT")
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = self.PRIMARY_COLOR
        title_p.paragraph_format.space_before = Pt(36)
        title_p.paragraph_format.space_after = Pt(12)

        sub_p = self.doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sub = sub_p.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  |  Mode: {mode}")
        run_sub.font.size = Pt(10)
        run_sub.font.color.rgb = self.SECONDARY_COLOR
        sub_p.paragraph_format.space_after = Pt(28)

        # Cover Page Matrix Table
        h2 = self.doc.add_heading("Traceability Matrix", level=1)
        h2.runs[0].font.color.rgb = self.PRIMARY_COLOR

        if matrix.headers and matrix.rows:
            table = self.doc.add_table(rows=len(matrix.rows) + 1, cols=len(matrix.headers))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False

            # Header Row
            hdr_cells = table.rows[0].cells
            for idx, text in enumerate(matrix.headers):
                hdr_cells[idx].text = text
                self._set_cell_background(hdr_cells[idx], "1E3A8A")
                for p in hdr_cells[idx].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(255, 255, 255)
                        r.font.size = Pt(10)

            # Data Rows
            for r_idx, row in enumerate(matrix.rows, start=1):
                row_cells = table.rows[r_idx].cells
                bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
                for c_idx, cell_data in enumerate(row):
                    self._set_cell_background(row_cells[c_idx], bg_color)
                    p = row_cells[c_idx].paragraphs[0]
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.space_before = Pt(2)

                    if cell_data.get("is_link") and cell_data.get("url"):
                        self._add_hyperlink(p, cell_data["url"], cell_data["text"])
                    else:
                        r = p.add_run(cell_data["text"])
                        r.font.size = Pt(9.5)

        self.doc.add_paragraph().paragraph_format.space_after = Pt(20)

        # ----------------------------------------------------------------------
        # STEP 2: INCOMPLETE ROWS & OPEN ISSUES AUDIT
        # ----------------------------------------------------------------------
        h_audit = self.doc.add_heading("Traceability Gap & Open Issues Audit", level=1)
        h_audit.runs[0].font.color.rgb = self.PRIMARY_COLOR

        incomplete_items = [item for item in audit_results if item["is_incomplete"]]

        if not incomplete_items:
            p = self.doc.add_paragraph()
            r = p.add_run("✓ All traceability matrix items are fully mapped and linked issues are closed.")
            r.font.bold = True
            r.font.color.rgb = RGBColor(21, 128, 61)
        else:
            p_warn = self.doc.add_paragraph()
            r_warn = p_warn.add_run(f"⚠ Found {len(incomplete_items)} incomplete row(s) in the Traceability Matrix:")
            r_warn.font.bold = True
            r_warn.font.color.rgb = self.ALERT_COLOR

            audit_table = self.doc.add_table(rows=len(incomplete_items) + 1, cols=4)
            audit_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_c = audit_table.rows[0].cells
            hdr_c[0].text = "Row #"
            hdr_c[1].text = "Requirement / Test Case"
            hdr_c[2].text = "Audit Findings / Status"
            hdr_c[3].text = "Linked Pull Request"

            for c in hdr_c:
                self._set_cell_background(c, "78350F")
                for p in c.paragraphs:
                    for r in p.runs:
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(255, 255, 255)
                        r.font.size = Pt(9.5)

            for i_idx, item in enumerate(incomplete_items, start=1):
                row_c = audit_table.rows[i_idx].cells
                self._set_cell_background(row_c[0], "FEF3C7")
                self._set_cell_background(row_c[1], "FEF3C7")
                self._set_cell_background(row_c[2], "FEF3C7")
                self._set_cell_background(row_c[3], "FEF3C7")

                row_c[0].paragraphs[0].text = str(item["row_index"])

                # Requirement / Test Case
                p_item = row_c[1].paragraphs[0]
                if item["req"].get("url"):
                    self._add_hyperlink(p_item, item["req"]["url"], item["req"]["text"])
                else:
                    p_item.text = item["req"]["text"]

                if item["tc"]["text"] != "--":
                    p_item.add_run(" / ")
                    if item["tc"].get("url"):
                        self._add_hyperlink(p_item, item["tc"]["url"], item["tc"]["text"])
                    else:
                        p_item.add_run(item["tc"]["text"])

                # Reasons
                p_reasons = row_c[2].paragraphs[0]
                p_reasons.text = "; ".join(item["reasons"])

                # PR links
                p_prs = row_c[3].paragraphs[0]
                if item["linked_prs"]:
                    for pr in item["linked_prs"]:
                        pr_num = pr["number"]
                        pr_url = pr.get("html_url") or f"https://github.com/mansikalraa/traceability-matrix-poc/pull/{pr_num}"
                        self._add_hyperlink(p_prs, pr_url, f"PR #{pr_num} ({pr['state']})")
                        p_prs.add_run(" ")
                else:
                    p_prs.text = "None"

        # Mode 1 stops here
        if mode == "1":
            self.doc.save(self.output_path)
            return

        # ----------------------------------------------------------------------
        # STEP 3 & STEP 4: STACKED DOCUMENTATION + PR REVIEWS + LINKAGE
        # ----------------------------------------------------------------------
        self.doc.add_page_break()
        h_docs = self.doc.add_heading("Stacked Specifications & Documentation", level=1)
        h_docs.runs[0].font.color.rgb = self.PRIMARY_COLOR

        for sec in stacked_docs:
            p_cat = self.doc.add_paragraph()
            r_cat = p_cat.add_run(f"[{sec['category'].upper()}]  {sec['file']}")
            r_cat.font.bold = True
            r_cat.font.size = Pt(11)
            r_cat.font.color.rgb = self.SECONDARY_COLOR
            p_cat.paragraph_format.space_before = Pt(16)
            p_cat.paragraph_format.space_after = Pt(4)

            parsed_elements = parse_md_elements(sec["path"])
            for elem in parsed_elements:
                etype = elem["type"]
                text = elem.get("text", "")

                if etype == "h1":
                    h = self.doc.add_heading(text, level=2)
                    h.runs[0].font.color.rgb = self.PRIMARY_COLOR
                elif etype == "h2":
                    h = self.doc.add_heading(text, level=3)
                    h.runs[0].font.color.rgb = self.SECONDARY_COLOR
                elif etype == "h3":
                    h = self.doc.add_heading(text, level=4)
                elif etype == "paragraph":
                    p = self.doc.add_paragraph(text)
                    p.paragraph_format.space_after = Pt(4)
                elif etype == "bullet":
                    p = self.doc.add_paragraph(text, style="List Bullet")
                    p.paragraph_format.space_after = Pt(2)
                elif etype == "numbered":
                    p = self.doc.add_paragraph(text, style="List Number")
                    p.paragraph_format.space_after = Pt(2)
                elif etype == "code":
                    p = self.doc.add_paragraph(text)
                    p.style = "Macro Text"
                elif etype == "spacer":
                    pass

            # Embed PR Details for matching file / section if applicable
            matched_prs = []
            file_code = sec["file"].replace(".md", "").upper()
            for item in audit_results:
                if file_code in item["tc"]["text"] or file_code in item["req"]["text"]:
                    matched_prs.extend(item["linked_prs"])

            # Also check PR list directly for matching title/branch
            for pr in gh_client.fetch_all_prs():
                if file_code.lower() in pr["title"].lower() or file_code.lower() in (pr.get("body") or "").lower():
                    if pr not in matched_prs:
                        matched_prs.append(pr)

            if matched_prs:
                for pr in matched_prs:
                    self._render_pr_details_box_docx(pr, gh_client)

            # Divider between stacked files
            p_hr = self.doc.add_paragraph()
            p_hr.paragraph_format.space_before = Pt(12)
            p_hr.paragraph_format.space_after = Pt(12)
            run_hr = p_hr.add_run("―" * 45)
            run_hr.font.color.rgb = self.SECONDARY_COLOR

        self.doc.save(self.output_path)

    def _render_pr_details_box_docx(self, pr, gh_client):
        """Embeds a styled PR Review Details callout box."""
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.rows[0].cells[0]
        self._set_cell_background(cell, "F0FDF4")  # Mint tint

        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)

        r_title = p.add_run("Latest Pull Request Review Details:\n")
        r_title.font.bold = True
        r_title.font.color.rgb = RGBColor(22, 101, 52)

        pr_num = pr["number"]
        pr_url = pr.get("html_url") or f"https://github.com/mansikalraa/traceability-matrix-poc/pull/{pr_num}"
        p.add_run("PR: ")
        self._add_hyperlink(p, pr_url, f"#{pr_num} — {pr['title']}")

        author = pr.get("user", {}).get("login", "Unknown")
        p.add_run(f"\nAuthor: {author}  |  State: {pr.get('state').upper()}")

        # Fetch reviews
        reviews = gh_client.fetch_pr_reviews(pr_num)
        reviewers = [r.get("user", {}).get("login") for r in reviews if r.get("user")]
        if not reviewers and pr.get("requested_reviewers"):
            reviewers = [r.get("login") for r in pr.get("requested_reviewers") if r.get("login")]

        rev_str = ", ".join(set(reviewers)) if reviewers else "No formal reviews recorded"
        p.add_run(f"\nReviewers: {rev_str}")

        pr_body = pr.get("body") or "No description provided."
        p.add_run(f"\nPR Message:\n{pr_body.strip()}")


# ==============================================================================
# PDF Document Builder (ReportLab)
# ==============================================================================

class NumberedCanvas(canvas.Canvas):
    """Two-pass canvas for dynamic total page count."""

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_header_footer(num_pages)
            super().showPage()
        super().save()

    def draw_header_footer(self, page_count):
        self.saveState()
        self.setFont("Helvetica", 9)
        self.setFillColor(colors.HexColor("#475569"))

        # Header (pages > 1)
        if self._pageNumber > 1:
            self.drawString(54, 11 * inch - 36, "Traceability Matrix & Compliance Report")
            self.setStrokeColor(colors.HexColor("#CBD5E1"))
            self.setLineWidth(0.5)
            self.line(54, 11 * inch - 42, 8.5 * inch - 54, 11 * inch - 42)

        # Footer
        page_str = f"Page {self._pageNumber} of {page_count}"
        self.drawRightString(8.5 * inch - 54, 36, page_str)
        self.drawString(54, 36, "CONFIDENTIAL — Internal Compliance Document")
        self.setStrokeColor(colors.HexColor("#CBD5E1"))
        self.setLineWidth(0.5)
        self.line(54, 48, 8.5 * inch - 54, 48)

        self.restoreState()


class PdfReportBuilder:
    """Builds formatted PDF documents using ReportLab."""

    def __init__(self, output_path="traceability_report.pdf"):
        self.output_path = output_path
        self.styles = getSampleStyleSheet()
        self._create_styles()

    def _create_styles(self):
        self.title_style = ParagraphStyle(
            "DocTitle",
            parent=self.styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=22,
            leading=26,
            textColor=colors.HexColor("#1E3A8A"),
            alignment=1,
            spaceAfter=12,
        )

        self.subtitle_style = ParagraphStyle(
            "DocSubTitle",
            parent=self.styles["Normal"],
            fontName="Helvetica",
            fontSize=10,
            leading=12,
            textColor=colors.HexColor("#475569"),
            alignment=1,
            spaceAfter=24,
        )

        self.h1_style = ParagraphStyle(
            "H1Heading",
            parent=self.styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=16,
            leading=20,
            textColor=colors.HexColor("#1E3A8A"),
            spaceBefore=16,
            spaceAfter=10,
            keepWithNext=True,
        )

        self.h2_style = ParagraphStyle(
            "H2Heading",
            parent=self.styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=13,
            leading=16,
            textColor=colors.HexColor("#334155"),
            spaceBefore=12,
            spaceAfter=6,
            keepWithNext=True,
        )

        self.body_style = ParagraphStyle(
            "DocBody",
            parent=self.styles["Normal"],
            fontName="Helvetica",
            fontSize=9.5,
            leading=13,
            textColor=colors.HexColor("#0F172A"),
            spaceAfter=6,
        )

        self.bullet_style = ParagraphStyle(
            "DocBullet",
            parent=self.body_style,
            leftIndent=15,
            bulletIndent=5,
            spaceAfter=3,
        )

        self.code_style = ParagraphStyle(
            "DocCode",
            parent=self.styles["Normal"],
            fontName="Courier",
            fontSize=8.5,
            leading=11,
            textColor=colors.HexColor("#1E293B"),
            backColor=colors.HexColor("#F1F5F9"),
            borderPadding=6,
            spaceAfter=6,
        )

    def build_report(self, matrix, audit_results, stacked_docs, gh_client, mode="2"):
        doc = SimpleDocTemplate(
            self.output_path,
            pagesize=letter,
            leftMargin=54,
            rightMargin=54,
            topMargin=54,
            bottomMargin=54,
        )

        story = []

        # ----------------------------------------------------------------------
        # STEP 1: COVER PAGE
        # ----------------------------------------------------------------------
        story.append(Spacer(1, 20))
        story.append(Paragraph("TRACEABILITY MATRIX &<br/>COMPLIANCE REPORT", self.title_style))
        story.append(Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  |  Mode: {mode}", self.subtitle_style))
        story.append(Spacer(1, 10))

        story.append(Paragraph("Traceability Matrix", self.h1_style))

        if matrix.headers and matrix.rows:
            table_data = []
            # Headers
            hdr_row = [Paragraph(f"<b>{h}</b>", ParagraphStyle("TH", parent=self.body_style, textColor=colors.white, alignment=1)) for h in matrix.headers]
            table_data.append(hdr_row)

            # Data
            for row in matrix.rows:
                r_cells = []
                for cell in row:
                    if cell.get("is_link") and cell.get("url"):
                        cell_p = Paragraph(f'<a href="{cell["url"]}" color="#1E40AF"><u>{cell["text"]}</u></a>', self.body_style)
                    else:
                        cell_p = Paragraph(cell["text"], self.body_style)
                    r_cells.append(cell_p)
                table_data.append(r_cells)

            t_style = TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E3A8A")),
                ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ])
            for r_idx in range(1, len(table_data)):
                bg = colors.HexColor("#F8FAFC") if r_idx % 2 == 1 else colors.white
                t_style.add("BACKGROUND", (0, r_idx), (-1, r_idx), bg)

            t = Table(table_data, colWidths=[1.2 * inch, 1.6 * inch, 1.4 * inch, 1.2 * inch, 1.1 * inch])
            t.setStyle(t_style)
            story.append(t)

        story.append(Spacer(1, 18))

        # ----------------------------------------------------------------------
        # STEP 2: INCOMPLETE ROWS & OPEN ISSUES AUDIT
        # ----------------------------------------------------------------------
        story.append(Paragraph("Traceability Gap & Open Issues Audit", self.h1_style))

        incomplete_items = [item for item in audit_results if item["is_incomplete"]]

        if not incomplete_items:
            story.append(Paragraph("<b>✓ All traceability matrix items are fully mapped and linked issues are closed.</b>", self.body_style))
        else:
            story.append(Paragraph(f'<font color="#B91C1C"><b>⚠ Found {len(incomplete_items)} incomplete row(s) in the Traceability Matrix:</b></font>', self.body_style))
            story.append(Spacer(1, 6))

            audit_table_data = []
            audit_hdr = [
                Paragraph("<b>Row #</b>", ParagraphStyle("TH1", parent=self.body_style, textColor=colors.white)),
                Paragraph("<b>Requirement / Test Case</b>", ParagraphStyle("TH2", parent=self.body_style, textColor=colors.white)),
                Paragraph("<b>Audit Findings / Status</b>", ParagraphStyle("TH3", parent=self.body_style, textColor=colors.white)),
                Paragraph("<b>Linked PR</b>", ParagraphStyle("TH4", parent=self.body_style, textColor=colors.white)),
            ]
            audit_table_data.append(audit_hdr)

            for item in incomplete_items:
                row_num_p = Paragraph(str(item["row_index"]), self.body_style)

                items_str = ""
                if item["req"].get("url"):
                    items_str += f'<a href="{item["req"]["url"]}" color="#1E40AF"><u>{item["req"]["text"]}</u></a>'
                else:
                    items_str += item["req"]["text"]

                if item["tc"]["text"] != "--":
                    items_str += " / "
                    if item["tc"].get("url"):
                        items_str += f'<a href="{item["tc"]["url"]}" color="#1E40AF"><u>{item["tc"]["text"]}</u></a>'
                    else:
                        items_str += item["tc"]["text"]
                item_p = Paragraph(items_str, self.body_style)

                reasons_p = Paragraph("; ".join(item["reasons"]), self.body_style)

                pr_links = []
                if item["linked_prs"]:
                    for pr in item["linked_prs"]:
                        pr_num = pr["number"]
                        pr_url = pr.get("html_url") or f"https://github.com/mansikalraa/traceability-matrix-poc/pull/{pr_num}"
                        pr_links.append(f'<a href="{pr_url}" color="#1E40AF"><u>PR #{pr_num} ({pr["state"]})</u></a>')
                prs_p = Paragraph(", ".join(pr_links) if pr_links else "None", self.body_style)

                audit_table_data.append([row_num_p, item_p, reasons_p, prs_p])

            at_style = TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#78350F")),
                ("BACKGROUND", (0, 1), (-1, -1), colors.HexColor("#FEF3C7")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#F59E0B")),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ])

            at = Table(audit_table_data, colWidths=[0.6 * inch, 1.8 * inch, 2.5 * inch, 1.6 * inch])
            at.setStyle(at_style)
            story.append(at)

        # Mode 1 stops here
        if mode == "1":
            doc.build(story, canvasmaker=NumberedCanvas)
            return

        # ----------------------------------------------------------------------
        # STEP 3 & STEP 4: STACKED DOCUMENTATION + PR REVIEWS + LINKAGE
        # ----------------------------------------------------------------------
        story.append(PageBreak())
        story.append(Paragraph("Stacked Specifications & Documentation", self.h1_style))

        for sec in stacked_docs:
            story.append(Spacer(1, 10))
            story.append(Paragraph(f"<b>[{sec['category'].upper()}] {sec['file']}</b>", self.h2_style))

            parsed_elements = parse_md_elements(sec["path"])
            for elem in parsed_elements:
                etype = elem["type"]
                text = elem.get("text", "")

                if etype == "h1":
                    story.append(Paragraph(text, self.h1_style))
                elif etype == "h2":
                    story.append(Paragraph(text, self.h2_style))
                elif etype == "h3":
                    story.append(Paragraph(text, ParagraphStyle("H3", parent=self.h2_style, fontSize=11, leading=14)))
                elif etype == "paragraph":
                    story.append(Paragraph(text, self.body_style))
                elif etype == "bullet":
                    story.append(Paragraph(f"• {text}", self.bullet_style))
                elif etype == "numbered":
                    story.append(Paragraph(text, self.bullet_style))
                elif etype == "code":
                    story.append(Paragraph(html.escape(text).replace("\n", "<br/>"), self.code_style))
                elif etype == "spacer":
                    story.append(Spacer(1, 4))

            # Embed PR Details Box
            matched_prs = []
            file_code = sec["file"].replace(".md", "").upper()
            for item in audit_results:
                if file_code in item["tc"]["text"] or file_code in item["req"]["text"]:
                    matched_prs.extend(item["linked_prs"])

            for pr in gh_client.fetch_all_prs():
                if file_code.lower() in pr["title"].lower() or file_code.lower() in (pr.get("body") or "").lower():
                    if pr not in matched_prs:
                        matched_prs.append(pr)

            if matched_prs:
                for pr in matched_prs:
                    self._render_pr_details_box_pdf(story, pr, gh_client)

            story.append(Spacer(1, 8))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#CBD5E1"), spaceAfter=10))

        doc.build(story, canvasmaker=NumberedCanvas)

    def _render_pr_details_box_pdf(self, story, pr, gh_client):
        pr_num = pr["number"]
        pr_url = pr.get("html_url") or f"https://github.com/mansikalraa/traceability-matrix-poc/pull/{pr_num}"

        reviews = gh_client.fetch_pr_reviews(pr_num)
        reviewers = [r.get("user", {}).get("login") for r in reviews if r.get("user")]
        if not reviewers and pr.get("requested_reviewers"):
            reviewers = [r.get("login") for r in pr.get("requested_reviewers") if r.get("login")]
        rev_str = ", ".join(set(reviewers)) if reviewers else "No formal reviews recorded"

        author = pr.get("user", {}).get("login", "Unknown")
        pr_body = html.escape(pr.get("body") or "No description provided.").replace("\n", "<br/>")

        box_text = f"""
        <font color="#166534"><b>Latest Pull Request Review Details:</b></font><br/>
        <b>PR:</b> <a href="{pr_url}" color="#1E40AF"><u>#{pr_num} — {pr['title']}</u></a><br/>
        <b>Author:</b> {author} &nbsp;|&nbsp; <b>State:</b> {pr.get('state').upper()}<br/>
        <b>Reviewers:</b> {rev_str}<br/>
        <b>PR Message:</b><br/>{pr_body}
        """

        p_box = Paragraph(box_text, ParagraphStyle("PRBox", parent=self.body_style, leading=13))
        t_box = Table([[p_box]], colWidths=[6.5 * inch])
        t_box.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F0FDF4")),
            ("BORDER", (0, 0), (-1, -1), 0.5, colors.HexColor("#86EFAC")),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
            ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ]))

        story.append(Spacer(1, 6))
        story.append(KeepTogether([t_box]))
        story.append(Spacer(1, 6))


# ==============================================================================
# CLI Entry Point
# ==============================================================================

def main():
    parser = argparse.ArgumentParser(
        description="Generate Traceability Matrix & Documentation (.docx / .pdf) in Mode 1 or Mode 2."
    )
    parser.add_argument(
        "--mode",
        choices=["1", "2"],
        default="2",
        help="Execution Mode: '1' for Steps 1-2 (Summary/Audit), '2' for Steps 1-4 (Full Report with Stacked Docs & PR Reviews). Default: 2.",
    )
    parser.add_argument(
        "--format",
        choices=["docx", "pdf", "both"],
        default="docx",
        help="Output document format: 'docx', 'pdf', or 'both'. Default: docx.",
    )
    parser.add_argument(
        "--output",
        default=None,
        help="Base path or file path for generated document.",
    )
    parser.add_argument(
        "--token",
        default=None,
        help="GitHub personal access token (optional, auto-detected if omitted).",
    )
    parser.add_argument(
        "--readme",
        default="README.md",
        help="Path to README.md containing the Traceability Matrix.",
    )
    parser.add_argument(
        "--docs-dir",
        default="docs",
        help="Path to directory containing documentation markdown files.",
    )

    args = parser.parse_args()

    print(f"=== Starting Traceability Report Generation (Mode: {args.mode}, Format: {args.format}) ===")

    # 1. Initialize GitHub API Client
    gh_client = GitHubClient(token=args.token)

    # 2. Parse Traceability Matrix
    print(f"Parsing Traceability Matrix from {args.readme}...")
    matrix = TraceabilityMatrix(readme_path=args.readme)
    print(f"Found {len(matrix.rows)} matrix row(s).")

    # 3. Fetch GitHub Issues & PRs
    print("Fetching repository issues and pull requests from GitHub...")
    issues = gh_client.fetch_all_issues()
    prs = gh_client.fetch_all_prs()
    issues_map = {i["number"]: i for i in issues}
    print(f"Fetched {len(issues)} issue(s) and {len(prs)} pull request(s).")

    # 4. Perform Matrix Audit
    audit_results = audit_matrix(matrix, issues_map, prs)
    incomplete_count = sum(1 for item in audit_results if item["is_incomplete"])
    print(f"Audit completed: {incomplete_count} incomplete row(s) flagged.")

    # 5. Load Stacked Markdown Documentation
    stacked_docs = load_markdown_files(docs_dir=args.docs_dir)
    print(f"Loaded {len(stacked_docs)} markdown documentation file(s) for stacking.")

    # 6. Build Document Outputs
    formats = ["docx", "pdf"] if args.format == "both" else [args.format]

    for fmt in formats:
        if args.output:
            out_path = args.output if args.output.endswith(f".{fmt}") else f"{args.output}.{fmt}"
        else:
            prefix = "summary" if args.mode == "1" else "traceability_report"
            out_path = f"{prefix}.{fmt}"

        print(f"Generating {fmt.upper()} report -> {out_path}...")

        if fmt == "docx":
            builder = DocxReportBuilder(output_path=out_path)
            builder.build_report(matrix, audit_results, stacked_docs, gh_client, mode=args.mode)
        elif fmt == "pdf":
            builder = PdfReportBuilder(output_path=out_path)
            builder.build_report(matrix, audit_results, stacked_docs, gh_client, mode=args.mode)

        print(f"✓ Successfully generated: {out_path}")

    print("=== Processing Completed Successfully ===")


if __name__ == "__main__":
    main()
